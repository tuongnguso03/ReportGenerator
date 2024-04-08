import tiktoken
import json
import docx
import io
from docx.shared import Inches

def json_split(string):
    data = json.loads(string.replace("'", '"'))
    ret = {
        "title" : data['report']['title'],
        "introduction" : data['report']['introduction'],
        "body" : data['report']['body'],
        "conclusion" : data['report']['conclusion'],
        "recommendations" : data['report']['recommendations']
    }
    return ret

def token_count(string):
    encoding = tiktoken.get_encoding("cl100k_base")
    """Returns the number of tokens in a text string."""
    num_tokens = len(encoding.encode(string))
    return num_tokens

def create_docx(title, columns_titles, columns_contents):
    doc = docx.Document()
    doc.add_heading(title, 0)
    for i in range(len(columns_titles)):
        doc.add_heading(columns_titles[i], 1)
        for paragraph in columns_contents[i].split("\n"):
            doc.add_paragraph(paragraph)

    bio = io.BytesIO()
    doc.save(bio)
    return bio

    

